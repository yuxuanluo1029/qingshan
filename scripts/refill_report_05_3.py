from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


BASE_DIR = Path(r"C:\temp\qsjd")
TEMPLATE_PATH = BASE_DIR / "submission_05-3_template.docx"
OUTPUT_PATH = BASE_DIR / "report_05-3_detailed.docx"
IMAGES_DIR = BASE_DIR / "images"


def set_para(doc: Document, index: int, text: str) -> None:
    doc.paragraphs[index - 1].text = text


def set_image(doc: Document, index: int, image_name: str, width_cm: float = 15.0) -> None:
    paragraph = doc.paragraphs[index - 1]
    paragraph.text = ""
    run = paragraph.add_run()
    run.add_picture(str(IMAGES_DIR / image_name), width=Cm(width_cm))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def fill_cover(doc: Document) -> None:
    set_para(doc, 6, "作品编号：QSJD-AI-2026-01")
    set_para(doc, 7, "作品名称：城迹——城市历史文化研学智能体")
    set_para(doc, 8, "")
    set_para(doc, 9, "填写日期：2026年04月02日")


def fill_chapter_1(doc: Document) -> None:
    set_para(
        doc,
        13,
        "“城迹”项目面向人工智能实践赛“智能教育与文化”方向，聚焦“历史+文化+城市”的复合场景，目标是把传统“城市打卡”升级为“可学习、可讲解、可复盘”的研学过程。项目来源于对原有文旅系统的实地观察：用户知道去哪里，却常常不知道为什么看、看什么、看完如何表达。为解决这一断层，城迹通过AI导师、研学路线、腾讯地图点位、现场观察任务和学习报告，将一次城市出行转化为一条完整的学习链路。"
    )
    set_para(doc, 14, "设计原理与创意来源")
    set_para(
        doc,
        15,
        "作品设计原理采用“目标驱动+证据驱动”双轨模式。目标驱动指系统先收集用户的城市、时长、兴趣主题和学习偏好，再生成与其能力相匹配的研学路线；证据驱动指每个点位都提供“历史文化简介、知识亮点、现场观察任务、互动提问、拍照与记录建议”，让用户在现场获得可验证的学习证据，而非只停留在感性印象。"
    )
    set_para(
        doc,
        16,
        "作品服务对象包括三类：一是中学与高校研学课程团队，二是博物馆与城市文化机构的公众教育活动组织者，三是希望进行深度城市探索的普通用户。核心价值是把“看城市”变成“读城市”：用户不仅能完成路线通行，还能形成讲解提纲、观察记录和学习反馈，便于课堂汇报、答辩展示与后续复盘。"
    )
    set_para(doc, 17, "应用页面示意")
    set_image(doc, 18, "01_home.png", 15.2)
    set_para(doc, 19, "图1  城迹首页（智能教育与文化方向入口）")
    set_para(doc, 20, "应用价值与推广前景")
    set_para(doc, 21, "表1  系统测试与交付指标")
    set_para(
        doc,
        22,
        "从应用价值看，城迹适合课程化使用：课前用于制定目标与路线，课中用于观察与问答，课后用于报告与反馈；从推广前景看，系统可平滑拓展到城市通识教育、文博社教活动、文旅公共服务等场景，且具备前后端分离部署能力，便于在公共网络快速上线。"
    )
    set_para(doc, 23, "")


def fill_chapter_2(doc: Document) -> None:
    set_para(doc, 25, "问题来源")
    set_para(
        doc,
        26,
        "城市研学正从“观光式参观”走向“任务化学习”，但市场上大量产品仍以导航或景点清单为主，缺乏教学目标映射和学习过程设计。团队在原项目迭代中发现，用户常出现三种典型问题：行程信息很多但知识结构混乱；现场体验很强但缺少可沉淀记录；回到课堂后难以形成高质量表达。"
    )
    set_para(
        doc,
        27,
        "此外，比赛答辩强调“方案设计+系统实现+功能演示+效果证据”一体化。如果系统只有内容展示而没有可验证数据，评审很难判断作品有效性。因此，本项目将接口联调结果、页面截图、结构化数据规模和部署路径纳入同一报告，形成可追溯证据链。"
    )
    set_para(doc, 28, "现有解决方案")
    set_para(
        doc,
        29,
        "现有方案主要有三类。第一类是地图导航类产品，优势是定位精准、路径规划成熟，短板是学习内容薄弱；第二类是文旅内容类平台，优势是图文视频丰富，但缺少任务化引导和个性化交互；第三类是通用问答机器人，优势是对话自然，但常与真实点位、路线和场景脱节。"
    )
    set_para(
        doc,
        30,
        "与上述方案相比，城迹的改进方向是把“路线、地图、对话、任务、报告”打通，构建“输入参数→生成方案→现场执行→结果反馈”的闭环，使产品既能在日常使用中提供教育价值，也能在比赛现场完整演示系统能力。"
    )
    set_para(doc, 31, "本作品要解决的痛点问题")
    set_para(
        doc,
        32,
        "痛点1：从“去哪”到“学什么”的目标断层；痛点2：从“看景”到“做观察”的方法断层；痛点3：从“现场体验”到“课后复盘”的成果断层；痛点4：从“技术实现”到“评审可证”的证据断层。"
    )
    set_para(
        doc,
        33,
        "围绕四类痛点，作品提出四项对应目标：路线必须可解释、节点必须可执行、问答必须可讲解、结果必须可复核。最终交付不仅是一个可运行系统，更是一套可用于教育场景的研学方法框架。"
    )
    set_para(doc, 34, "")
    set_para(doc, 35, "解决问题的思路")
    set_para(
        doc,
        36,
        "功能需求：支持城市、时长、兴趣主题、学习偏好输入；支持AI导师对话、路线建议、知识问答；支持腾讯地图点位标注与信息弹窗；支持节点任务执行与研学报告生成。性能需求：关键接口返回稳定JSON结构，前端可构建、后端可部署，具备课堂演示时的连续可用性。"
    )
    set_para(
        doc,
        37,
        "数据层采用结构化TypeScript对象，集中于frontend/src/data/mockData.ts，当前真实规模为：5位导师角色、5座城市（杭州/成都/长沙/西安/武汉）、25个文化点位、15条研学路线模板、25条任务模板（由点位映射生成）、5份城市学习报告模板。点位字段包含经纬度、类型、知识亮点、观察任务、拍照建议与推荐理由，可直接支撑地图与教学联动。"
    )
    set_para(
        doc,
        38,
        "评价指标在第5章验证，包括：接口健康状态（/health）、核心API可用性（/api/ai/chat、/api/ai/knowledge、/api/ai/route）、页面流程完整性（设置→对话→路线→节点→报告）、工程构建可用性（npm run build）以及部署可复现性（Vercel+Render）。"
    )


def fill_chapter_3(doc: Document) -> None:
    set_para(
        doc,
        40,
        "技术方案采用“前端交互层+后端服务层+模型与地图能力层”三层架构。前端基于React+Vite+TypeScript，负责参数采集、页面路由和可视化呈现；后端基于Node.js+Express，提供统一AI网关与业务接口；能力层接入腾讯地图GL与WebService能力，以及大模型推理能力。模型网关支持AI_PROVIDER=aliyun/tencent/auto三种模式：默认阿里云百炼（Qwen），保留腾讯混元SDK接入能力，实现密钥后置和弹性切换。"
    )
    set_image(doc, 41, "02_setup.png", 14.6)
    set_para(doc, 42, "图2  研学设置页（城市、时长、文化主题与学习偏好输入）")


def fill_chapter_4(doc: Document) -> None:
    set_para(
        doc,
        44,
        "系统实现坚持“基于现有工程重构”的原则，不推翻原有React+Vite结构，而是对页面语义、数据模型与交互逻辑做定向升级。页面流程为：Index说明作品定位；Setup收集研学参数；Chat提供导师式对话；Route展示路线与地图；NodeDetail提供任务执行指引；Report输出学习画像。核心改造文件覆盖mockData.ts、ai-service.ts、TencentMap.tsx、tencent-lbs.ts及各页面组件，确保视觉风格延续、主题叙事彻底去电竞化。"
    )
    set_image(doc, 45, "04_route.png", 14.6)
    set_para(doc, 46, "图3  研学路线页（地图点位、学习目标、推荐理由与顺序）")


def fill_chapter_5(doc: Document) -> None:
    set_para(
        doc,
        48,
        "测试采用“接口层—流程层—工程层”三级验证。测试环境为Windows 10、Node.js v24.14.1、npm 11.11.0。本地联调中，GET /health返回success=true、providerMode=aliyun、model=aliyun:qwen-plus；核心接口POST /api/ai/chat、POST /api/ai/route、POST /api/ai/knowledge均可正常返回。一次实测延迟为chat 1477ms、route 1539ms、knowledge 2112ms。前端执行npm run build成功，生成可部署dist产物，满足比赛现场演示与复现要求。"
    )
    set_image(doc, 49, "05_node.png", 14.2)
    set_para(doc, 50, "图4  节点详情页（知识亮点、观察任务、互动提问与地图定位）")

    table = doc.tables[0]
    rows = [
        ["测试维度", "结果A（接口健康）", "结果B（功能联调）", "结果C（工程构建）"],
        ["验证项1", "GET /health 返回 success=true", "/api/ai/chat、/api/ai/route、/api/ai/knowledge 全部可用", "npm run build 成功"],
        ["验证项2", "providerMode=aliyun", "设置→对话→路线→节点→报告流程可完整走通", "dist 产物生成且可部署"],
        ["验证项3", "model=aliyun:qwen-plus", "chat 1477ms / route 1539ms / knowledge 2112ms", "前后端分离部署配置已补齐"],
        ["验证项4", "providers: aliyun=ready, tencent=missing_key", "腾讯地图点位加载与弹窗交互正常", "支持 Vercel + Render 公网部署方案"],
        ["结论", "服务状态可观测", "功能满足答辩演示要求", "具备稳定交付与复现能力"],
    ]
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value


def fill_chapter_6_and_refs(doc: Document) -> None:
    set_para(
        doc,
        52,
        "本作品在创意、工程与验证三个层面形成闭环：创意上完成从“电竞文旅”到“历史文化教育”的方向重构；工程上补齐前后端能力链路与地图联动；验证上提供接口实测、流程演示和构建结果。整体工作量覆盖主题重定位、数据重写、提示词重构、页面重构、后端补建、文档补齐和部署配置，能够支撑现场答辩“讲得清、演得出、证得实”。"
    )
    set_para(doc, 53, "作品特色与创新点")
    set_para(
        doc,
        54,
        "创新点一：将城市空间转化为任务化学习现场，建立“参数设置→路线生成→节点观察→学习反馈”闭环。创新点二：把路线规划、地图定位、智能问答和学习报告统一在同一系统内，避免“工具割裂”。创新点三：服务端模型网关支持阿里云与腾讯云模式切换，既保障部署稳定性，也便于后续能力扩展。创新点四：保留国风沉浸视觉并转化叙事语义，使界面风格服务于教育目标。"
    )
    set_para(doc, 55, "应用推广")
    set_para(
        doc,
        56,
        "在教育场景中，城迹可作为“城市历史文化研学课”数字底座：教师预设学习目标，学生现场完成观察任务，课后自动形成报告。"
    )
    set_para(
        doc,
        57,
        "在文博场景中，城迹可作为社教活动的互动讲解工具：将馆内展陈与城市空间联动，帮助公众理解“展品—场景—时代”的关系。部署方面采用前端Vercel+后端Render的轻量方案，也可迁移至腾讯云静态托管+云函数架构，具备较好的推广弹性。"
    )
    set_para(doc, 58, "作品展望")
    set_para(doc, 59, "展望1：扩展更多城市与学段版本，形成小学、初中、高中、大学差异化任务模板。")
    set_para(doc, 60, "展望2：新增教师端编排与班级管理功能，支持任务发布、批注反馈与过程性评价。")
    set_para(doc, 61, "展望3：接入更多开放文博数据与活动日历，提高路线推荐的时效性和情境适配度。")
    set_para(doc, 62, "展望4：建立学习前测/后测与行为日志指标体系，用量化数据进一步验证教育效果。")
    set_para(doc, 64, "[1] 城迹项目组. README.md, 2026.")
    set_para(doc, 65, "[2] 城迹项目组. docs/方案设计.md, 2026.")
    set_para(
        doc,
        66,
        "[3] 城迹项目组. docs/用户手册.md, 2026. [4] 腾讯位置服务. 腾讯地图 JavaScript GL API 与 WebService 文档. [5] 阿里云百炼（DashScope）兼容接口文档. [6] 腾讯云混元大模型 Node.js SDK 文档."
    )


def main() -> None:
    doc = Document(str(TEMPLATE_PATH))
    fill_cover(doc)
    fill_chapter_1(doc)
    fill_chapter_2(doc)
    fill_chapter_3(doc)
    fill_chapter_4(doc)
    fill_chapter_5(doc)
    fill_chapter_6_and_refs(doc)
    doc.save(str(OUTPUT_PATH))
    print(f"saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
